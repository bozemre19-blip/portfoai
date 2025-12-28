#!/usr/bin/env python3
"""
Development Report Generator
Fills the development report template with AI-generated content based on child observations.
"""

import sys
import json
from datetime import datetime
from docx import Document

def fill_report_template(template_path, data, output_path):
    """
    Fill the development report template with provided data.
    
    Args:
        template_path: Path to the template DOCX file
        data: Dictionary containing all report data
        output_path: Path where the filled report will be saved
    """
    try:
        # Load template
        doc = Document(template_path)
        
        # Table 1: Basic Information
        table1 = doc.tables[0]
        table1.rows[0].cells[1].text = data.get('childName', '')
        table1.rows[0].cells[3].text = data.get('birthDate', '')
        table1.rows[1].cells[1].text = data.get('teacherName', '')
        table1.rows[1].cells[3].text = data.get('schoolStartDate', '')
        table1.rows[2].cells[1].text = data.get('reportDate', datetime.now().strftime('%d.%m.%Y'))
        
        # Table 2: Skill Areas
        table2 = doc.tables[1]
        table2.rows[2].cells[0].text = data.get('alanBecerileri', '')
        table2.rows[4].cells[0].text = data.get('sosyalDuygusal', '')
        table2.rows[6].cells[0].text = data.get('kavramsal', '')
        table2.rows[8].cells[0].text = data.get('okuryazarlik', '')
        table2.rows[10].cells[0].text = data.get('degerler', '')
        table2.rows[12].cells[0].text = data.get('egilimler', '')
        table2.rows[14].cells[0].text = data.get('genelDegerlendirme', '')
        
        # Save filled document
        doc.save(output_path)
        return True
        
    except Exception as e:
        print(f"Error filling template: {str(e)}", file=sys.stderr)
        return False

def main():
    """Main entry point for CLI usage"""
    if len(sys.argv) < 4:
        print("Usage: python generate_report.py <template_path> <data_json> <output_path>")
        sys.exit(1)
    
    template_path = sys.argv[1]
    data_json = sys.argv[2]
    output_path = sys.argv[3]
    
    # Parse data
    try:
        data = json.loads(data_json)
    except json.JSONDecodeError as e:
        print(f"Invalid JSON data: {e}", file=sys.stderr)
        sys.exit(1)
    
    # Generate report
    success = fill_report_template(template_path, data, output_path)
    
    if success:
        print(json.dumps({"success": True, "outputPath": output_path}))
    else:
        print(json.dumps({"success": False, "error": "Failed to generate report"}))
        sys.exit(1)

if __name__ == "__main__":
    main()
